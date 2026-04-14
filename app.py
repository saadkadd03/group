# app.py
from flask import Flask, render_template, request, redirect, url_for, send_file
import sqlite3
import json
import random
import io
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from docx import Document

app = Flask(__name__)
DATABASE = "submissions.db"
SECRET_CODE = "admin123"   # ← Change this

# ================== DATA ==================
STUDENTS = [
    "Saad Kaddioui", "Hamza Iaaza", "Tarik Laatabi" , "Fatima Ezzahra Follouss", "Hamza Bagrou",
    "ABDERRAHMAN EL GHAZY", "Yahya Sabri", "Safa Charhbili", "Hajar Amakha",
    "Marwa El Jayyi", "Kaoutar Abaabous", "Fadna El mikh", "Laila Ait-Mouloud",
    "Sanae El Khayyati" , "Imane Lkemil", "Imane ALAHYANE",
    "Asma Fikri", "Aissa Khalek", "Ahmed Boukhir", "Omar Ouboussag",
    "Essamlali Jihane", "AISSAM ID BELLAL", "AHMED ESSAIDI", "Khalid Moukhlis",
    "Mohamed Khriss", "SOUAD NIDBRKA", "Elhoucine Frimane", "Malika Lhaloui",
    "Hanane Hanane Naama", "Nour-Elhouda Mansouri", "Hamza Ajrar", "Hajar Akherraz",
    "Imane Barhdadi", "Khaoula KANTAOUI", "Khaoula Radi", "Hanane Abaraghe" , "Jihane Essamlali"
]

PRESENTATIONS = [
    {"title": "Testing, Assessment, Evaluation, Quiz, Exam, Measurement & Teaching", "size": 1},
    {"title": "Requirements of Good Tests", "size": 2},
    {"title": "Formative vs Summative Assessment", "size": 2},
    {"title": "Formal vs Informal Assessment", "size": 2},
    {"title": "Criterion vs Norm Referenced Testing", "size": 2},
    {"title": "High vs Low Stakes Tests", "size": 2},
    {"title": "Performance-Based Assessment", "size": 2},
    {"title": "Stages of Test Design", "size": 2},
    {"title": "Formative Assessment (Laura Greenstein)", "size": 3},
    {"title": "MCQ Construction", "size": 2},
    {"title": "Alternative Assessment", "size": 3},
    {"title": "Assessing Grammar & Vocabulary", "size": 2},
    {"title": "Assessing Reading & Listening", "size": 2},
    {"title": "Assessing Writing", "size": 2},
    {"title": "Assessing Speaking", "size": 2},
    {"title": "Washback Effect", "size": 2},
    {"title": "Moroccan High School Assessment", "size": 2},
    {"title": "Cheating", "size": 2}
]

# ================== DB SETUP ==================
def init_db():
    conn = sqlite3.connect(DATABASE)
    conn.execute('''
        CREATE TABLE IF NOT EXISTS submissions (
            student_name TEXT PRIMARY KEY,
            preferred_teammates TEXT,
            preferred_presentations TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.execute('''
        CREATE TABLE IF NOT EXISTS current_assignment (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            groups TEXT NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

# ================== STRONGER GROUPING LOGIC ==================
def generate_assignments(submissions):
    try:
        from pulp import LpProblem, LpMaximize, LpVariable, lpSum, value, LpStatus, PULP_CBC_CMD
    except ImportError:
        print("PuLP not installed. Using fallback.")
        return generate_assignments_fallback(submissions)

    if not submissions:
        return generate_assignments_fallback(submissions)

    student_data = {sub['student']: sub for sub in submissions}
    all_students = STUDENTS[:]

    prob = LpProblem("UniPresent_Group_Assignment", LpMaximize)

    x = {(s, g_idx): LpVariable(f"x_{s}_{g_idx}", cat='Binary')
         for s in all_students for g_idx in range(len(PRESENTATIONS))}

    satisfaction = 0

    for s in all_students:
        data = student_data.get(s, {})
        pref_topics = data.get('presentations', [])
        pref_teammates = data.get('teammates', [])

        for g_idx, group in enumerate(PRESENTATIONS):
            var = x[(s, g_idx)]

            # Topic preference (very important)
            if group['title'] in pref_topics:
                satisfaction += 30 * var

            # STRONGER Teammate preference (main fix)
            for t in pref_teammates:
                if t in all_students:
                    t_var = x.get((t, g_idx))
                    if t_var:
                        satisfaction += 18 * var * t_var   # Increased weight

    prob += satisfaction

    # Constraints
    for s in all_students:
        prob += lpSum(x[(s, g_idx)] for g_idx in range(len(PRESENTATIONS))) == 1

    for g_idx, group in enumerate(PRESENTATIONS):
        prob += lpSum(x[(s, g_idx)] for s in all_students) == group['size']

    # Solve
    status = prob.solve(PULP_CBC_CMD(msg=False))

    if LpStatus[status] != 'Optimal':
        print("Solver warning → using fallback")
        return generate_assignments_fallback(submissions)

    # Build groups
    groups = [dict(p, assigned=[]) for p in PRESENTATIONS]
    for s in all_students:
        for g_idx, group in enumerate(PRESENTATIONS):
            if value(x[(s, g_idx)]) > 0.5:
                groups[g_idx]['assigned'].append(s)
                break

    for g in groups:
        g['assigned'].sort()

    return groups


def generate_assignments_fallback(submissions):
    students_list = STUDENTS[:]
    random.shuffle(students_list)
    groups = [dict(p, assigned=[]) for p in PRESENTATIONS]
    idx = 0
    for group in groups:
        while len(group['assigned']) < group['size'] and idx < len(students_list):
            group['assigned'].append(students_list[idx])
            idx += 1
    return groups


# ================== PDF & WORD (unchanged) ==================
def create_pdf(groups):
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.setFont("Helvetica-Bold", 24)
    c.drawString(1.5*inch, 10*inch, "UniPresent - Final Group Assignments")
    c.setFont("Helvetica", 11)
    c.drawString(1.5*inch, 9.6*inch, f"Generated on: {datetime.now().strftime('%B %d, %Y at %H:%M')}")
    y = 9*inch
    for group in groups:
        if not group.get('assigned'): continue
        c.setFont("Helvetica-Bold", 13)
        c.drawString(1.5*inch, y, group['title'])
        y -= 22
        c.setFont("Helvetica", 11)
        c.drawString(1.8*inch, y, f"Assigned ({len(group['assigned'])}/{group['size']}): {', '.join(group['assigned'])}")
        y -= 35
        if y < 1.5*inch:
            c.showPage()
            y = 10*inch
    c.save()
    buffer.seek(0)
    return buffer


def create_word_doc(groups):
    doc = Document()
    doc.add_heading('UniPresent - Final Group Assignments', 0)
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
    for group in groups:
        if not group.get('assigned'): continue
        doc.add_heading(group['title'], level=2)
        p = doc.add_paragraph()
        p.add_run(f'Assigned: {len(group["assigned"])} / {group["size"]} students').bold = True
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No.'
        hdr_cells[1].text = 'Student Name'
        for i, student in enumerate(group['assigned'], 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = student
        doc.add_paragraph('')
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ================== ROUTES (same as before) ==================
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        student_name = request.form.get('student_name')
        teammates = request.form.getlist('teammates')
        presentations = request.form.getlist('presentations')
        if student_name in teammates:
            teammates.remove(student_name)

        conn = sqlite3.connect(DATABASE)
        conn.execute('''
            INSERT OR REPLACE INTO submissions 
            (student_name, preferred_teammates, preferred_presentations)
            VALUES (?, ?, ?)
        ''', (student_name, json.dumps(teammates), json.dumps(presentations)))
        conn.commit()
        conn.close()
        return redirect(url_for('index'))

    return render_template('index.html', students=STUDENTS,
                           presentations=[p['title'] for p in PRESENTATIONS])


@app.route('/admin', methods=['GET', 'POST'])
def admin():
    error = None
    if request.method == 'POST':
        if request.form.get('code') == SECRET_CODE:
            conn = sqlite3.connect(DATABASE)
            conn.row_factory = sqlite3.Row
            rows = conn.execute("SELECT * FROM submissions").fetchall()
            conn.close()

            submissions = [
                {
                    'student': row['student_name'],
                    'teammates': json.loads(row['preferred_teammates']) if row['preferred_teammates'] else [],
                    'presentations': json.loads(row['preferred_presentations']) if row['preferred_presentations'] else []
                }
                for row in rows
            ]

            groups = generate_assignments(submissions)

            conn = sqlite3.connect(DATABASE)
            conn.execute("DELETE FROM current_assignment")
            conn.execute("INSERT INTO current_assignment (id, groups) VALUES (1, ?)", (json.dumps(groups),))
            conn.commit()
            conn.close()
            return redirect(url_for('results'))
        else:
            error = "❌ Invalid admin code!"

    return render_template('admin.html', error=error)


@app.route('/results')
def results():
    conn = sqlite3.connect(DATABASE)
    row = conn.execute("SELECT groups FROM current_assignment WHERE id=1").fetchone()
    conn.close()
    if not row or not row[0]:
        return "No assignment generated yet. Go to Admin first.", 404
    groups = json.loads(row[0])
    return render_template('results.html', groups=groups)


@app.route('/download_pdf')
def download_pdf():
    conn = sqlite3.connect(DATABASE)
    row = conn.execute("SELECT groups FROM current_assignment WHERE id=1").fetchone()
    conn.close()
    if not row or not row[0]:
        return "No assignment found.", 404
    groups = json.loads(row[0])
    return send_file(create_pdf(groups), as_attachment=True,
                     download_name='UniPresent_Group_Assignments.pdf',
                     mimetype='application/pdf')


@app.route('/download_word')
def download_word():
    conn = sqlite3.connect(DATABASE)
    row = conn.execute("SELECT groups FROM current_assignment WHERE id=1").fetchone()
    conn.close()
    if not row or not row[0]:
        return "No assignment found.", 404
    groups = json.loads(row[0])
    return send_file(create_word_doc(groups), as_attachment=True,
                     download_name='UniPresent_Group_Assignments.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


if __name__ == "__main__":
    init_db()
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
